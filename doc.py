from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import json
'''
Doc為產生word報告格式
report_word 為測試報告的第一頁 ,file_name為檔案名稱,dev_name為設備名稱,test_case_array為每個測項所需的值
api_doc為測試api測試報告詳細測試過程
ui_doc為測試前端測試報告詳細測試過程
'''
class Doc():
    def report_word(self,file_name, dev_name, test_case_array = []):
        self.document = Document()
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        title = self.document.add_paragraph()
        tle =title.add_run('思納捷科技')
        tle.bold =True
        tle.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nowtime  = datetime.date.today().strftime("%Y/%m/%d")
        records = [
            ('臺北市105民生東路四段133號3F之9', '日期:'+str(nowtime)),
            ('電話(02) 7713-4828  傳真 (02) 2713-0867', '頁數:共'+str(len(test_case_array)+1)+'頁')
        ]

        table = self.document.add_table(rows= 0,cols= 2)
        for t1,t2 in records: 
            tab = table.add_row()
            row_cells=tab.cells
            row_cells[0].text = t1
            row_cells[0].paragraphs[0].paragraph_format.space_after  = Pt(0)
            row_cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)
            row_cells[1].text = t2
            row_cells[1].paragraphs[0].paragraph_format.space_after  = Pt(0)
            row_cells[1].paragraphs[0].paragraph_format.space_before = Pt(0)
            row_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        test_report = self.document.add_paragraph()
        trp = test_report.add_run('測試報告')
        trp.bold =True
        trp.underline = True
        trp.font.size = Pt(12)
        test_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
        test_report.paragraph_format.space_after  = Pt(0)
        test_report.paragraph_format.space_before = Pt(0)

        dev_id = self.document.add_paragraph()
        devid = dev_id.add_run('設備名稱:       ')
        devid.font.size = Pt(12)
        devid1 = dev_id.add_run(dev_name)
        devid1.font.color.rgb = RGBColor(255, 0, 0)
        devid1.italic = True
        devid1.font.size = Pt(12)
        dev_id.paragraph_format.space_after  = Pt(0)
        dev_id.paragraph_format.space_before = Pt(0)
        
        test_time = self.document.add_paragraph()
        ttime = test_time.add_run('測試時間:       ')
        ttime.font.size=Pt(12)
        ttime1 = test_time.add_run(str(nowtime)+"~"+str(nowtime)+"\n")
        ttime1.font.size = Pt(11)
        test_time.paragraph_format.space_after  = Pt(0)
        test_time.paragraph_format.space_before = Pt(0)

        self.test_project = self.document.add_table(rows= 1 , cols= 3)
        self.test_project.style = 'Table Grid'
        hdr_cells = self.test_project.rows[0].cells
        hdr_cells[0].text = "Test Case"
        hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[0].width = Inches(20)
        hdr_cells[1].text = "Target"
        hdr_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[1].width = Inches(20)
        hdr_cells[2].text = "Result"
        hdr_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].width = Inches(1)
        if test_case_array!=[]:
            if len(test_case_array[0]) == 9:
                self.api_doc(test_case_array)
                self.document.save('.\\api_report\\'+file_name+".docx")
            elif len(test_case_array[0]) == 5:
                self.ui_doc(test_case_array)
                self.document.save('.\\api_ui\\'+file_name+".docx")
            else:
                return False
        else:
            return False
        return True
    
    def api_doc(self,test_case_array):
        #第一頁的表格
        count = 1
        for i in test_case_array:
            tc = i[0]
            tg = i[1]
            rt = i[2]
            tp = self.test_project.add_row().cells
            if rt == "pass":
                color = '00ff00'
            else:
                color = 'ff0000'
            ttc = tp[2]._tc.get_or_add_tcPr() #拿到Result格
            #設定Result格子顏色
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'pct100')   #設定100%填充
            shd.set(qn('w:fill'), color)    #設定顏色
            tp[0].text = "test"+str(count)+"\n"+tc
            tp[1].text = tg
            tp[2].text = rt
            tp[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            count += 1
            ttc.append(shd)
        self.document.add_page_break()

        #每個測試結果
        num=1
        for j in test_case_array:
            tc1 = j[0]
            tg1 = j[1]
            rt1 = j[2]
            step1 = j[3]
            method1 = j[4]
            url1 = j[5]
            data1 = j[6]
            rsp1 = j[7]
            except_rt1 = j[8]

            test_table = self.document.add_table(rows= 6 , cols= 2)
            test_table.style = 'Table Grid'
            table_name = test_table.rows[0].cells
            table_name[0].text = "測項名稱"
            table_name[0].width = Inches(1.5)
            table_name[0]._tc.get_or_add_tcPr()
            shd1 = OxmlElement('w:shd')
            shd1.set(qn('w:val'),'pct100')
            shd1.set(qn('w:fill'),'F9CD88') 
            table_name[0]._tc.get_or_add_tcPr().append(shd1)

            table_name[1].text = tc1
            table_name[1].paragraphs[0].runs[0].font.size = Pt(12)
            table_name[1].width = Inches(8.5)
            table_name[1]._tc.get_or_add_tcPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'),'pct100')
            shd2.set(qn('w:fill'),'F9CD88') 
            table_name[1]._tc.get_or_add_tcPr().append(shd2)

            table_target = test_table.rows[1].cells
            table_target[0].text = "測試目的"
            table_target[1].text = tg1
            table_target[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_step = test_table.rows[2].cells
            table_step[0].text = "測試步驟"
            table_step[1].text = str(step1).replace(",","\n").translate({ord("'"):None}).strip("[]")
            table_step[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_condition = test_table.rows[3].cells
            table_condition[0].text = "測試條件"
            table_condition[1].text = tg1
            table_condition[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_reslut = test_table.rows[4].cells
            table_reslut[0].text = "測試結果"
            table_reslut[1].text = rt1
            table_reslut[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_content = test_table.rows[5].cells
            table_content[0].text = "測試內容"
            table_content[1].text = "Req:\nmethod:"+method1+"\nUrl:"+url1+"\ndata:"+str(json.dumps(data1,indent=4, ensure_ascii=False).encode('utf8').decode("utf8"))+"\n------------------------------------------------\nRsp:\n"+str(json.dumps(rsp1,indent=4, ensure_ascii=False).encode('utf8').decode("utf8"))+"\n------------------------------------------------\n預期結果:\n"+str(json.dumps(except_rt1,indent=4, ensure_ascii=False).encode('utf8').decode("utf8"))
            table_content[1].paragraphs[0].runs[0].font.size = Pt(12)

            if num < len(test_case_array):
                self.document.add_page_break()
                num+=1 

    def ui_doc(self,test_case_array):
        #第一頁的表格
        count = 1
        for i in test_case_array:
            tc = i[0]
            tg = i[1]
            rt = i[2]
            tp = self.test_project.add_row().cells
            if rt == "pass":
                color = '00ff00'
            else:
                color = 'ff0000'
            ttc = tp[2]._tc.get_or_add_tcPr() #拿到Result格
            #設定Result格子顏色
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'pct100')   #設定100%填充
            shd.set(qn('w:fill'), color)    #設定顏色
            tp[0].text = "test"+str(count)+"\n"+tc
            tp[1].text = tg
            tp[2].text = rt
            tp[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            count += 1
            ttc.append(shd)
        self.document.add_page_break()

        #每個測試結果
        num=1
        for j in test_case_array:
            tc1 = j[0]
            tg1 = j[1]
            rt1 = j[2]
            step1 = j[3]
            pricture1 = j[4] 
        
            test_table = self.document.add_table(rows= 6 , cols= 2)
            test_table.style = 'Table Grid'
            table_name = test_table.rows[0].cells
            table_name[0].text = "測項名稱"
            table_name[0].width = Inches(2)
            table_name[0]._tc.get_or_add_tcPr()
            shd1 = OxmlElement('w:shd')
            shd1.set(qn('w:val'),'pct100')
            shd1.set(qn('w:fill'),'F9CD88') 
            table_name[0]._tc.get_or_add_tcPr().append(shd1)

            table_name[1].text = tc1
            table_name[1].paragraphs[0].runs[0].font.size = Pt(12)
            table_name[1].width = Inches(7.5)
            table_name[1]._tc.get_or_add_tcPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'),'pct100')
            shd2.set(qn('w:fill'),'F9CD88') 
            table_name[1]._tc.get_or_add_tcPr().append(shd2)

            table_target = test_table.rows[1].cells
            table_target[0].text = "測試目的"
            table_target[1].text = tg1
            table_target[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_step = test_table.rows[2].cells
            table_step[0].text = "測試步驟"
            table_step[1].text = str(step1).replace(",","\n").translate({ord("'"):None}).strip("[]")
            table_step[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_condition = test_table.rows[3].cells
            table_condition[0].text = "測試條件"
            table_condition[1].text = tg1
            table_condition[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_reslut = test_table.rows[4].cells
            table_reslut[0].text = "測試結果"
            table_reslut[1].text = rt1
            table_reslut[1].paragraphs[0].runs[0].font.size = Pt(12)

            table_content = test_table.rows[5].cells
            table_content[0].text = "測試內容"
            table_content[1].text ="\n"
            for i in pricture1:
                table_content[1].paragraphs[0].add_run().add_picture("D:\\Demo\\測試\\screenshot\\"+str(i),width=Inches(4.5),height=Inches(2.5))  
                table_content[1].paragraphs[0].add_run("\n")

            if num < len(test_case_array):
                self.document.add_page_break()
                num+=1